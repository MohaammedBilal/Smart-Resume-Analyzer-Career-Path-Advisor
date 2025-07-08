# import necessary libraries
import os
import fitz  # PyMuPDF
import docx
import re
import spacy
from sentence_transformers import SentenceTransformer, util
import torch
import numpy as np
import joblib
import time
import requests
from fpdf import FPDF
from bs4 import BeautifulSoup
from sklearn.metrics.pairwise import cosine_similarity

# Ensure the necessary models and data files are available
# Load small English NLP model
# !python -m spacy download en_core_web_sm
nlp = spacy.load("en_core_web_sm")
model = SentenceTransformer('all-MiniLM-L6-v2')  # Fast and good for semantic similarity

# Load pre-trained semantic model and role embeddings
semantic_model_name = joblib.load("semantic_model_name.pkl")
semantic_model = SentenceTransformer('all-MiniLM-L6-v2')
role_labels = joblib.load("semantic_roles.pkl")
role_embeddings = np.load("semantic_role_embeddings.npy")


# Function to load keywords from a file
def load_keywords_from_file(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return [line.strip().lower() for line in f if line.strip()]

#  Clean Unicode characters in text    
def clean_unicode(text):
    return (text.replace("–", "-")
                .replace("—", "-")
                .replace("’", "'")
                .replace("“", '"')
                .replace("”", '"')
                .replace("•", "-")
                .replace("…", "...")
                .encode("latin-1", "ignore").decode("latin-1"))


# Sample skill list (extend as needed)
SKILLS_DB = load_keywords_from_file("skills.txt")

# Read PDF files
def read_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

# Read DOCX files
def read_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# Extract text from uploaded files
def extract_text(file):
    if hasattr(file, "name") and file.name.endswith(".pdf"):
        return read_pdf(file)
    elif hasattr(file, "name") and file.name.endswith(".docx"):
        return read_docx(file)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX.")

# Clean and format text
def clean_text(text):
    return "\n".join([line.strip() for line in text.splitlines() if line.strip() != ""])

# Extract basic information from text

# Extract email, phone, and name from text
def extract_email(text):
    match = re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    return match.group(0) if match else None

def extract_phone(text):
    match = re.search(r"\+?\d[\d\s\-]{8,}\d", text)
    return match.group(0) if match else None

def extract_name(text):
    lines = text.splitlines()[:10]
    doc = nlp(" ".join(lines))
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return None

# Extract education, experience, skills, certifications, projects, and objective sections    
# Extract education section
def extract_education(text):
    education_keywords = load_keywords_from_file("education_keywords.txt")
    lines = text.lower().splitlines()
    education = []
    in_education_section = False
    for line in lines:
        lower_line = line.lower()
        if any(keyword in lower_line for keyword in education_keywords):
            education.append(line.strip())
            in_education_section = True
        elif in_education_section and line.strip() and not any(keyword in lower_line for keyword in ["experience", "skills", "certifications", "projects", "objective"]):
            education.append(line.strip())
        else:
            in_education_section = False
    return education

# Extract experience section
def extract_experience(text):
    lines = text.lower().splitlines()
    experience = []
    in_experience_section = False
    for line in lines:
        lower_line = line.lower()
        if "experience" in lower_line or "work history" in lower_line:
            in_experience_section = True
            continue
        if in_experience_section and line.strip() and not any(keyword in lower_line for keyword in ["education", "skills", "certifications", "projects", "objective"]):
            experience.append(line.strip())
        else:
            in_experience_section = False
    return experience 

    
# Extract section based on keywords and stop conditions
def extract_section(text, section_names, stop_keywords=[]):
    lines = text.splitlines()
    section_content = []
    in_section = False
    for line in lines:
        lower_line = line.lower()
        if any(name in lower_line for name in section_names):
            in_section = True
            continue
        if in_section:
            if any(keyword in lower_line for keyword in stop_keywords):
                break
            if line.strip():
                section_content.append(line.strip())
    return section_content

# === Certification Extraction Functions ===
def extract_certifications(text):
    cert_keywords = load_keywords_from_file("cert_keywords.txt")
    stop_keywords = ["skills", "experience", "education", "projects", "objective"]
    return extract_section(text, cert_keywords, stop_keywords)


# === Projects Extraction Functions ===    
def extract_projects(text):
    project_keywords = load_keywords_from_file("project_keywords.txt")
    stop_keywords = ["skills", "experience", "education", "certifications", "objective"]
    return extract_section(text, project_keywords, stop_keywords)

# === Objective Extraction Functions ===
def extract_objective(text):
    objective_keywords = load_keywords_from_file("objective_keywords.txt")
    stop_keywords = ["education", "experience", "skills", "certifications", "projects"]
    content = extract_section(text, objective_keywords, stop_keywords)
    return content[:5]

#  === Skills Extraction Functions ===
def extract_skills(text):
    text_lower = text.lower()
    skills_found = [skill for skill in SKILLS_DB if skill in text_lower]
    return list(set(skills_found))

def extract_skills_from_text(text, skills_db):
    text = text.lower()
    return list({skill for skill in skills_db if skill in text})

# === Final Resume Score Calculation ===
def final_resume_score(field_score, skill_score, sbert_score):
    return round((0.2 * field_score + 0.4 * skill_score + 0.4 * sbert_score), 2)    

"""# Live coursera scrapper"""
                            # ====== Live Course Recommendations from Coursera and Udemy ======

# This function scrapes Coursera for courses related to a specific skill.
def get_courses_from_coursera(skill, max_results=3):
    url = f"https://www.coursera.org/search?query={skill}"
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        course_elements = soup.select('li[data-e2e="SearchResult"]')
        courses = []
        for elem in course_elements[:max_results]:
            title_elem = elem.select_one('h3')
            link_elem = elem.select_one('a[href]')
            if title_elem and link_elem:
                title = title_elem.text.strip()
                link = "https://www.coursera.org" + link_elem['href']
                courses.append((title, link))
        if not courses:
            raise ValueError("No courses found")
        return courses
    except Exception as e:
        return [(f"'{skill}' courses on Coursera", f"https://www.coursera.org/search?query={skill}")]

# This function recommends courses based on missing skills by scraping Coursera.
def recommend_courses_live(missing_skills):
    recommendations = {}
    for skill in missing_skills:
        recommendations[skill] = get_courses_from_coursera(skill)
        time.sleep(1)  # Avoid overloading the server
    return recommendations


# This function scrapes Udemy for courses related to a specific skill.
def get_courses_from_udemy(skill, max_results=3):
    url = f"https://www.udemy.com/courses/search/?q={skill}"
    headers = {"User-Agent": "Mozilla/5.0"}
    try:
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, 'html.parser')
        course_elements = soup.select('div.udlite-focus-visible-target.udlite-heading-md.course-card--course-title--2f7tE')
        courses = []
        for elem in course_elements[:max_results]:
            title = elem.text.strip()
            link = "https://www.udemy.com" + elem.find_parent('a')['href']
            courses.append((title, link))
        if not courses:
            raise ValueError("No courses found")
        return courses
    except Exception as e:
        return [(f"'{skill}' courses on Udemy", f"https://www.udemy.com/courses/search/?q={skill}")]

# This function combines course recommendations from Coursera and Udemy.
def recommend_courses_combined(missing_skills):
    recommendations = {}
    for skill in missing_skills:
        coursera = get_courses_from_coursera(skill)
        udemy = get_courses_from_udemy(skill)
        recommendations[skill] = {"coursera": coursera, "udemy": udemy}
        time.sleep(1)  # avoid IP blocking
    return recommendations

# === Add Career Suggestions to PDF Report ===
def add_career_suggestions_to_pdf(pdf, recommendations):
    pdf.add_section("Career Path Suggestions")
    
    # Match standard text styling
    pdf.set_font("Arial", '', 11)
    pdf.set_text_color(0, 0, 0)  # Set text color to black
    
    for title, score in recommendations:
        clean_title = clean_unicode(str(title))
        pdf.multi_cell(0, 8, f"- {clean_title}: {score}% match")


ESSENTIAL_SECTIONS = ["objective", "projects", "certifications", "skills", "education", "experience"]

# === Career Path Prediction (Live Scraping from RemoteOK — Skill-Focused) ===
def get_job_recommendations_from_remoteok(resume_text, top_skills, max_jobs=5):
    job_titles, job_texts = [], []
    seen_jobs = set()
    headers = {"User-Agent": "Mozilla/5.0"}
    # model = SentenceTransformer('all-MiniLM-L6-v2')

    for keyword in top_skills:
        url = f"https://remoteok.com/remote-{keyword}-jobs"
        try:
            response = requests.get(url, headers=headers, timeout=10)
            soup = BeautifulSoup(response.text, 'html.parser')
            jobs = soup.select('tr.job')
            for job in jobs[:3]:
                title_tag = job.select_one("td.company.position h2")
                desc_tag = job.select_one("td.tags")
                title = title_tag.text.strip() if title_tag else "Unknown Title"
                description = desc_tag.text.strip() if desc_tag else "No description."
                job_key = title + description[:30]
                if job_key not in seen_jobs:
                    seen_jobs.add(job_key)
                    job_titles.append(title)
                    job_texts.append(description)
        except Exception as e:
            pass

    job_embeddings = model.encode(job_texts, convert_to_tensor=True)
    resume_embedding = model.encode(resume_text, convert_to_tensor=True)
    similarities = util.cos_sim(resume_embedding, job_embeddings)[0]
    scored_roles = list(zip(job_titles, similarities.cpu().tolist()))
    scored_roles.sort(key=lambda x: x[1], reverse=True)
    return [(title, round(score * 100, 2)) for title, score in scored_roles[:max_jobs]]

# === Section Aliases Loader ===
def load_section_aliases(filepath):
    section_aliases = {}
    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            if ':' in line:
                section, alias = line.strip().split(':', 1)
                section = section.strip().lower()
                alias = alias.strip().lower()
                section_aliases.setdefault(section, []).append(alias)
    return section_aliases

# === Semantic Role Prediction ===
# This function predicts future roles based on semantic similarity to skill-based role descriptions.
def predict_semantic_roles(user_skills, top_k=5):
    """
    Predict future roles based on semantic similarity to skill-based role descriptions.
    Args:
        user_skills (list): List of skill strings.
        top_k (int): Number of top roles to return.
    Returns:
        List of tuples: [(role, similarity_score), ...]
    """
    if not user_skills:
        return [("No skills detected", 0.0)]

    skill_text = ", ".join(user_skills).lower()
    input_embedding = semantic_model.encode(skill_text, convert_to_tensor=True)

    similarities = util.cos_sim(input_embedding, torch.tensor(role_embeddings))[0]
    top_indices = torch.topk(similarities, k=min(top_k, len(role_labels))).indices.tolist()

    return [(role_labels[i], round(similarities[i].item() * 100, 2)) for i in top_indices]



# === Resume Quality Assessment ===
# This function assesses the completeness of a resume based on the presence of essential sections.
def assess_resume_quality(text):
    """
    Assess resume completeness based on presence of essential sections.

    Args:
        text (str): Raw resume text

    Returns:
        tuple: (completeness_score, missing_sections)
            - completeness_score (float): 0-100% score
            - missing_sections (list): Missing essential sections
    """
    import re
    # Optional: from rapidfuzz import fuzz

    # 1. Extract potential section headers
    section_headers = set()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        clean_line = re.sub(r'[:\-\•\—]+$', '', line.lower()).strip()
        if not clean_line or len(clean_line.split()) > 14 :
            continue
        section_headers.add(clean_line)

    # 2. Load section aliases (with error handling)
    try:
        section_aliases = load_section_aliases("section_aliases.txt")
    except Exception as e:
        print(f"Error loading aliases: {e}")
        section_aliases = {sec: [sec] for sec in ESSENTIAL_SECTIONS}

    # 3. Match against essential sections
    found_sections = set()
    for section in ESSENTIAL_SECTIONS:
        aliases = set(section_aliases.get(section, [section]))
        aliases.add(section)  # Always include the canonical name

        for header in section_headers:
            # Exact match
            if any(alias == header for alias in aliases):
                found_sections.add(section)
                break
            # Substring match
            if any(alias in header for alias in aliases):
                found_sections.add(section)
                break
            # Optional: Fuzzy match (uncomment if rapidfuzz is installed)
            # if any(fuzz.ratio(alias, header) > 80 for alias in aliases):
            #     found_sections.add(section)
            #     break

    # 4. Calculate metrics
    missing = [sec for sec in ESSENTIAL_SECTIONS if sec not in found_sections]
    completeness = round((1 - len(missing)/len(ESSENTIAL_SECTIONS)) * 100, 1)
    return completeness, missing



# === PDF Report Class ===
# This class generates a PDF report for resume evaluation.
class PDF(FPDF):
    def header(self):
        self.set_font("Arial", 'B', 14)
        self.set_text_color(255, 255, 255)
        self.set_fill_color(40, 40, 40)
        self.cell(0, 10, "Resume Evaluation Report", ln=True, align='C', fill=True)
        self.ln(5)

    def add_section(self, title):
        self.set_font("Arial", 'B', 12)
        self.set_text_color(255, 255, 255)
        self.set_fill_color(60, 60, 60)
        self.cell(0, 10, title, ln=True, fill=True)
        self.ln(2)

    def add_list(self, items):
        self.set_font("Arial", '', 11)
        self.set_text_color(0, 0, 0)  # Darker text for better contrast on white background
        for item in items:
            clean_item = clean_unicode(str(item))
            self.multi_cell(0, 8, f"- {clean_item}")
        self.ln(1)

    def add_table(self, headers, data):
        self.set_font("Arial", 'B', 11)
        self.set_fill_color(200, 200, 200) 
        self.set_text_color(0, 0, 0)  
        for header in headers:
            self.cell(48, 8, clean_unicode(str(header)), border=1, align='C', fill=True)
        self.ln()
        self.set_font("Arial", '', 11)
        self.set_text_color(50, 50, 50) 
        for row in data:
            for item in row:
                self.cell(48, 8, str(item), border=1, align='C')
            self.ln()

# This function generates a PDF report for the resume evaluation for candidates.
def generate_pdf_report(name, email, phone, education, experience, skills,
                        matched_skills, missing_skills, field_score,
                        match_score, semantic_match_score, final_score,
                        recommended_courses, recommendations,
                        resume_text,
                        output_path="resume_report.pdf"):

    pdf = PDF()
    pdf.add_page()

    pdf.add_section("Basic Information")
    pdf.add_list([
        f"Name: {name}",
        f"Email: {email}",
        f"Phone: {phone}"
    ])

    pdf.add_section("Education")
    pdf.add_list(education or ["Not Found"])

    pdf.add_section("Experience")
    pdf.add_list(experience or ["Not Found"])

    pdf.add_section("Extracted Skills")
    pdf.add_list(skills or ["Not Found"])

    pdf.add_section("Scores Summary")
    pdf.add_table(
        ["Field Score", "Skill Score", "Semantic Score", "Final Score"],
        [[f"{field_score}%", f"{match_score}%", f"{semantic_match_score}%", f"{final_score}%"]]
    )

    pdf.add_section("Matched Skills")
    pdf.add_list(matched_skills or ["None"])

    pdf.add_section("Missing Skills")
    pdf.add_list(missing_skills or ["None"])

    pdf.add_section("Recommended Courses")
    for skill, platforms in recommended_courses.items():
        pdf.set_font("Arial", 'B', 11)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(0, 8, f"{skill.title()}:", ln=True)

        pdf.set_font("Arial", 'I', 11)
        pdf.set_text_color(30, 30, 120)
        pdf.cell(0, 8, "Coursera:", ln=True)

        pdf.set_font("Arial", '', 10)
        pdf.set_text_color(0, 0, 0)
        for title, link in platforms["coursera"]:
            pdf.multi_cell(0, 7, clean_unicode(f"- {title}: {link}"))

        pdf.set_font("Arial", 'I', 11)
        pdf.set_text_color(30, 30, 120)
        pdf.cell(0, 8, "Udemy:", ln=True)

        pdf.set_font("Arial", '', 10)
        pdf.set_text_color(0, 0, 0)
        for title, link in platforms["udemy"]:
            pdf.multi_cell(0, 7, clean_unicode(f"- {title}: {link}"))

        pdf.ln(2)

    pdf.add_section("AI Career Path Predictions (SBERT Model)")
    pdf.set_font("Arial", '', 11)
    pdf.set_text_color(0, 0, 0)
    for title, score in recommendations:
        clean_title = clean_unicode(str(title))
        pdf.multi_cell(0, 8, f"- {clean_title}: {score}% match")

    # Resume completeness
    section_data = {
        "objective": extract_objective(resume_text),
        "projects": extract_projects(resume_text),
        "certifications": extract_certifications(resume_text),
        "skills": skills,
        "education": education,
        "experience": experience
    }
    missing_sections = [sec for sec, content in section_data.items() if not content]
    quality_score = round((1 - len(missing_sections) / 6) * 100, 2)

    pdf.add_section("Resume Quality Assessment")
    pdf.add_list([
        f"Completeness Score: {quality_score}%",
        "Missing Sections:" if missing_sections else "All essential sections found."
    ] + missing_sections)

    pdf.output(output_path)


"""##  HR Phase"""

# This function evaluates a resume against multiple job descriptions (JDs) and returns scores.
def evaluate_resume_against_multiple_jds(resume_text, skills_db, jd_list):
    # model = SentenceTransformer('all-MiniLM-L6-v2')
    resume_embedding = model.encode(resume_text, convert_to_tensor=True)

    results = []
    for title, jd_text in jd_list:
        jd_embedding = model.encode(jd_text, convert_to_tensor=True)
        cosine_score = util.cos_sim(resume_embedding, jd_embedding)[0][0].item()
        semantic_match_score = round(cosine_score * 100, 2)

        jd_skills = [skill for skill in skills_db if skill.lower() in jd_text.lower()]
        resume_skills = [skill for skill in skills_db if skill.lower() in resume_text.lower()]

        matched_skills = list(set(jd_skills) & set(resume_skills))
        missing_skills = list(set(jd_skills) - set(resume_skills))
        match_score = round((len(matched_skills) / len(jd_skills)) * 100, 2) if jd_skills else 0

        final_score = round((0.2 * 100 + 0.4 * match_score + 0.4 * semantic_match_score), 2)

        results.append({
            "title": title,
            "match_score": match_score,
            "semantic_score": semantic_match_score,
            "final_score": final_score,
            "matched_skills": matched_skills,
            "missing_skills": missing_skills
        })
    return results

"""### For HR"""

# This function evaluates multiple resumes against multiple job descriptions (JDs) and returns scores.
def evaluate_multiple_resumes_against_jds(resume_texts_dict, skills_db, jd_list):
    # model = SentenceTransformer('all-MiniLM-L6-v2')
    results = []
    for resume_name, resume_text in resume_texts_dict.items():
        resume_embedding = model.encode(resume_text, convert_to_tensor=True)
        for title, jd_text in jd_list:
            jd_embedding = model.encode(jd_text, convert_to_tensor=True)
            cosine_score = util.cos_sim(resume_embedding, jd_embedding)[0][0].item()
            semantic_match_score = round(cosine_score * 100, 2)

            jd_skills = [skill for skill in skills_db if skill.lower() in jd_text.lower()]
            resume_skills = [skill for skill in skills_db if skill.lower() in resume_text.lower()]

            matched_skills = list(set(jd_skills) & set(resume_skills))
            missing_skills = list(set(jd_skills) - set(resume_skills))
            match_score = round((len(matched_skills) / len(jd_skills)) * 100, 2) if jd_skills else 0

            final_score = round((0.2 * 100 + 0.4 * match_score + 0.4 * semantic_match_score), 2)

            results.append({
                "resume": resume_name,
                "jd_title": title,
                "match_score": match_score,
                "semantic_score": semantic_match_score,
                "final_score": final_score
            })
    return results

# This function generates a PDF report for the resume evaluation for hr.
def generate_hr_summary_pdf(results, output_path="hr_resume_scores.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "HR Resume Evaluation Summary", ln=True, align='C')
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 11)
    pdf.cell(60, 8, "Resume Name", border=1, align='C')
    pdf.cell(60, 8, "Job Role", border=1, align='C')
    pdf.cell(30, 8, "Skill Match", border=1, align='C')
    pdf.cell(30, 8, "Final Score", border=1, align='C')
    pdf.ln()

    pdf.set_font("Arial", '', 10)
    for res in results:
        pdf.cell(60, 8, res["resume"], border=1)
        pdf.cell(60, 8, res["jd_title"], border=1)
        pdf.cell(30, 8, f"{res['match_score']}%", border=1, align='C')
        pdf.cell(30, 8, f"{res['final_score']}%", border=1, align='C')
        pdf.ln()
    pdf.output(output_path)
